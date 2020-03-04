#coding:utf-8

import docx
import glob, os
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt
from docx.oxml.ns import qn



path_now = os.getcwd()
path_captures = path_now + "\captures"

pic_height = 3880000

#可接受的圖片格式
valid_images = [".jpg",".JPG",".png",".PNG"]

#圖片路徑list
image_file_name = [fn for fn in os.listdir(path_captures)
              if any(fn.endswith(ext) for ext in valid_images)]

#照時間排序一下
image_file_name.sort()

image_file_name_noext =[]

for filenames in image_file_name:
    image_file_name_noext.append(filenames.split('.')[0])

image_file_path = []
for fn in image_file_name:
    image_file_path.append(os.path.join(path_now,"captures",fn))
    




document = docx.Document('template(dont-touch-it).docx')

#document = docx.Document()

#窄邊界
section = document.sections[0]
section.left_margin=Cm(1.27)
section.right_margin=Cm(1.27)
section.top_margin=Cm(1.27)
section.bottom_margin=Cm(1.27)

#TITLE
p = document.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_format = p.paragraph_format
p_format.line_spacing = Pt(15)
run = p.add_run(u"title")

font = run.font
font.bold = True
font.name= 'New Times Roman'   #設定英文字體
font.size=Pt(12)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')  #設定中文字體

p.bold = True




#先把表格一次畫完
tbl = document.add_table(rows=len(image_file_path)*2, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'

	

for r in range(0,len(image_file_path)):
    cell = tbl.cell(r*2, 0)
    pic = image_file_path[r]
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(pic, height = pic_height)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for r in range(1,len(image_file_path)+1):
    cell = tbl.cell(r*2-1, 0)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("%s" %(image_file_name_noext[r-1]))
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

document.save("output.docx")
